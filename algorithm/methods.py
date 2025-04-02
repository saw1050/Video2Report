# coding utf-8
'''
    Automatic video-to-reporting tool.
    -Version: 1.0
    -Briefly:
        Identify video content by whisper-large.
        Capture the keyframes that need to be represented by comparing significant changes between successive frames.
'''
import cv2
from PIL import Image, ImageDraw, ImageFont
from moviepy import VideoFileClip
import numpy as np
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from pathlib import Path
import shutil
from config import config
from . import model

cfg = config['current']
transcribers = model.Transcribers(model = cfg.MODEL, devices = cfg.DEVICES)

def create_report(video_file, doc_title, report_id, image_dump_path, username, threshold=30.0, min_interval=1.0):
    if Path(image_dump_path).exists():
        shutil.rmtree(image_dump_path)
    Path(image_dump_path).mkdir(parents=True, exist_ok=True)
    cap = cv2.VideoCapture(video_file)
    fps = cap.get(cv2.CAP_PROP_FPS)
    # num_frames = cap.get(cv2.CAP_PROP_FRAME_COUNT)
    frame_interval = int(fps * min_interval)
    prev_frame = None
    last_frame = None
    keyframes = []
    frame_count = 0
    while True:
        ret, frame = cap.read()
        if not ret:
            break
        frame_count += 1
        # progress[username]['report']['capture'] = frame_count/num_frames
        if prev_frame is None:
            prev_frame = frame
            continue
        gray_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        prev_gray_frame = cv2.cvtColor(prev_frame, cv2.COLOR_BGR2GRAY)
        diff = cv2.absdiff(prev_gray_frame, gray_frame)
        non_zero_count = np.count_nonzero(diff)
        diff_percentage = (non_zero_count / diff.size) * 100
        if diff_percentage > threshold and (len(keyframes) == 0 or (frame_count - keyframes[-1][0]) > frame_interval):
            keyframes.append((frame_count, prev_frame))
            last_frame = frame
        prev_frame = frame
    keyframes.append((frame_count, last_frame))
    cap.release()
    
    for idx, (frame_count, frame) in enumerate(keyframes):
        output_path = f"{image_dump_path}/{idx + 1}.jpg"
        cv2.imwrite(output_path, frame)
        # print(f"Saved keyframe {idx + 1} at frame {frame_count}({frame_count/fps}s) to {output_path}")

    transcriber = transcribers.get_transcriber()
    print(f'Start the transcription, remain devices: {len(transcribers)}')
    # progress[username]['report']['transcribe'] = 1
    result = transcriber.transcribe(video_file, fp16="False")
    transcribers.return_transcriber(transcriber)
    # progress[username]['report']['transcribe'] = 2
    print(f'End the transcription, remain devices: {len(transcribers)}')
    
    remain_segments = result['segments']
    key_idxes = [0]
    for idx, (frame_count, frame) in enumerate(keyframes):
        # progress[username]['report']['match'] = idx/len(keyframes)
        time_point = frame_count/fps
        min_span = 9999999
        for i, segment in enumerate(remain_segments):
            span = abs(segment['start'] - time_point)
            if span < min_span:
                min_span = span
            else:
                key_idxes.append(key_idxes[-1]+i-1)
                remain_segments = remain_segments[i-1:]
                break

    output_doc = Document()
    output_doc.styles['Normal'].font.name = u'宋体'
    output_doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    output_doc.add_heading(doc_title)
    count_pic = 0
    for i in range(len(key_idxes))[1:]:
        paragraph = ''
        for segment in result['segments'][key_idxes[i-1]:key_idxes[i]]:
            paragraph += segment['text'] + '\n'
        output_doc.add_paragraph(str(paragraph)[:-1])
        count_pic += 1
        output_doc.add_picture(f'{image_dump_path}/{count_pic}.jpg', Cm(12.))
        # progress[username]['report']['write'] = count_pic/len(key_idxes)
    
    paragraph = ''
    for segment in result['segments'][key_idxes[-1]:]:
        paragraph += segment['text'] + '\n'
    output_doc.add_paragraph(str(paragraph)[:-1])
    count_pic += 1
    output_doc.add_picture(f'{image_dump_path}/{count_pic}.jpg', Cm(12.))
    # progress[username]['report']['write'] = count_pic/len(key_idxes)
    
    doc_save_path = f'{cfg.DATA_DUMP_PATH}/{report_id}.docx'
    output_doc.save(doc_save_path)
    shutil.rmtree(image_dump_path)
    Path(video_file).unlink(missing_ok=True)
    
    return Path(doc_save_path).absolute()


def add_subtitles(video_file, font_size):
    video_format = video_file.split('.')[-1]
    output_file_path_without_shuffix = video_file.replace(f'.{video_format}', '')
    # result = transcriber.transcribe(video_file)
    transcriber = transcribers.get_transcriber()
    print(f'Start the transcription, remain devices: {len(transcribers)}')
    result = transcriber.transcribe(video_file, fp16="False")
    transcribers.return_transcriber(transcriber)
    print(f'End the transcription, remain devices: {len(transcribers)}')
    # subtitles_segments = result['segments']
    subtitles_segments = [{'start':segment['start'], 'end':segment['end'], 'text':segment['text']} for segment in result['segments']]
    cursor = 0
    
    audio_data = VideoFileClip(video_file).audio
    cap = cv2.VideoCapture(video_file)

    frame_width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    frame_height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
    fps = cap.get(cv2.CAP_PROP_FPS)

    fourcc = cv2.VideoWriter_fourcc(*'mp4v')
    out = cv2.VideoWriter(f'{output_file_path_without_shuffix}-mute.mp4', fourcc, fps, (frame_width, frame_height))

    font = ImageFont.truetype('ukai.ttc', font_size)
    font_color = (255, 255, 255)
    stroke_width = font_size//16
    stroke_color = (0, 0, 0)
    count_frame = 0
    while cap.isOpened():
        ret, frame = cap.read()
        if not ret:
            break
        count_frame += 1
        
        timepoint = count_frame/fps
        subtitles_segment = subtitles_segments[cursor]
        if timepoint >= round(subtitles_segment['end']) and cursor < len(subtitles_segments) - 1:
            cursor += 1
        if timepoint >= int(subtitles_segment['start']):
            frame_array = Image.fromarray(frame)
            pil_draw = ImageDraw.Draw(frame_array)
            
            subtitle_text = subtitles_segment['text'].replace('。', '')
            subtitle_len = len(subtitle_text)*font_size
            if subtitle_len > frame_width*0.8:
                split_line = int(len(subtitle_text)*0.6)
                subtitle_text_1, subtitle_text_2 = subtitle_text[:split_line], subtitle_text[split_line:]
                subtitle_len_1 = len(subtitle_text_1)*font_size
                subtitle_len_2 = len(subtitle_text_2)*font_size
                text_position_1 = (frame_width//2 - subtitle_len_1//2, frame_height - 2*font_size - int(font_size*0.7))
                text_position_2 = (frame_width//2 - subtitle_len_2//2, frame_height - 2*font_size + int(font_size*0.7))
                pil_draw.text(text_position_1, subtitle_text_1, fill=font_color, font=font, stroke_width=stroke_width, stroke_fill=stroke_color)
                pil_draw.text(text_position_2, subtitle_text_2, fill=font_color, font=font, stroke_width=stroke_width, stroke_fill=stroke_color)
            else:
                text_position = (frame_width//2 - subtitle_len//2, frame_height - 2*font_size)
                pil_draw.text(text_position, subtitle_text, fill=font_color, font=font, stroke_width=stroke_width, stroke_fill=stroke_color)

            frame = np.array(frame_array)
        out.write(frame)
    cap.release()
    out.release()
    cv2.destroyAllWindows()
    final_video = VideoFileClip(f'{output_file_path_without_shuffix}-mute.mp4')
    final_video.audio = audio_data
    final_video.write_videofile(f'{output_file_path_without_shuffix}-with-subtitles.mp4')
    
    Path(video_file).unlink(missing_ok=True)
    Path(f'{output_file_path_without_shuffix}-mute.mp4').unlink(missing_ok=True)
    
    return Path(f'{output_file_path_without_shuffix}-with-subtitles.mp4').absolute()